"""SPDX-FileCopyrightText: Ian Suhih
SPDX-License-Identifier: GPL-3.0-or-later
"""

from __future__ import annotations

import os
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont, TTFError
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from lopa_engine import draw_lopa_graph


def _normalize_text(value: Optional[str]) -> str:
    return (value or "").strip()


def _format_number(value) -> str:
    if value is None:
        return ""
    try:
        num = float(value)
    except (TypeError, ValueError):
        return str(value)
    if num == 0:
        return "0"
    abs_num = abs(num)
    if abs_num >= 1:
        return f"{num:.1f}"
    formatted = f"{num:.2g}"
    if "e" in formatted or "E" in formatted:
        formatted = f"{num:.6f}".rstrip("0").rstrip(".")
    return formatted


def _escape_pdf_text(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _split_bold_runs(value: str) -> List[Tuple[str, bool]]:
    parts = []
    segments = value.split("**")
    for idx, segment in enumerate(segments):
        if segment == "":
            continue
        parts.append((segment, idx % 2 == 1))
    return parts


def _parse_markdown_lines(text: str) -> List[Tuple[Optional[int], str]]:
    lines = (text or "").splitlines()
    parsed = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            parsed.append((None, ""))
            continue
        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            if 1 <= hashes <= 6 and stripped[hashes:hashes + 1] == " ":
                parsed.append((hashes, stripped[hashes + 1:]))
                continue
        parsed.append((0, line))
    return parsed


def _group_summaries(groups: List[Dict]) -> List[Dict]:
    summaries = []
    for group in groups:
        questions = group.get("questions", [])
        total_score = 0.0
        total_max = 0.0
        for q in questions:
            score = q.get("score")
            max_score = q.get("max_score")
            if score is not None:
                total_score += float(score)
            if max_score is not None:
                total_max += float(max_score)
        normalized = None
        if total_max > 0:
            normalized = total_score / total_max
        summaries.append(
            {
                "group_id": group.get("group_id"),
                "group_name": group.get("group_name"),
                "weight": group.get("weight"),
                "total_score": total_score,
                "total_max": total_max,
                "normalized": normalized,
                "questions": questions,
            }
        )
    return summaries


def _maturity_level(score: float) -> str:
    if score < 25:
        return "Critical"
    if score < 50:
        return "Low"
    if score < 75:
        return "Medium"
    return "High"


def _format_attacker(attacker_type: str, attacker_potential: str) -> str:
    return f"{attacker_type.capitalize()} attacker, capability: {attacker_potential.capitalize()}"


def _build_lopa_images(layers_before: List[Dict], layers_after: List[Dict], output_dir: str) -> Tuple[Optional[str], Optional[str]]:
    before_path = None
    after_path = None
    if layers_before:
        _, before_path = draw_lopa_graph(layers_before, os.path.join(output_dir, "lopa_before"))
    if layers_after:
        _, after_path = draw_lopa_graph(layers_after, os.path.join(output_dir, "lopa_after"))
    return before_path, after_path


def _register_pdf_font() -> str:
    candidates = [
        (
            "DejaVuSans",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ),
        (
            "NotoSans",
            "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        ),
        (
            "LiberationSans",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        ),
    ]
    for font_name, font_path, bold_path in candidates:
        if not os.path.exists(font_path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            bold_font_name = font_name
            if os.path.exists(bold_path):
                bold_font_name = f"{font_name}-Bold"
                pdfmetrics.registerFont(TTFont(bold_font_name, bold_path))
            pdfmetrics.registerFontFamily(font_name, normal=font_name, bold=bold_font_name)
            return font_name
        except (OSError, TTFError):
            pass
    return "Helvetica"


def _build_pdf(
    output_path: str,
    *,
    title: str,
    attacker_description: str,
    technical_groups: List[Dict],
    organizational_groups: List[Dict],
    layers_input: List[Dict],
    risk_result: Dict,
    lopa_before_path: Optional[str],
    lopa_after_path: Optional[str],
    recommendations_text: str,
) -> None:
    styles = getSampleStyleSheet()
    pdf_font = _register_pdf_font()
    heading1 = ParagraphStyle("Heading1", parent=styles["Heading1"], fontName=pdf_font, spaceAfter=8)
    heading2 = ParagraphStyle("Heading2", parent=styles["Heading2"], fontName=pdf_font, spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=pdf_font, leading=14)
    small = ParagraphStyle("Small", parent=styles["BodyText"], fontName=pdf_font, fontSize=9, leading=11)
    table_cell = ParagraphStyle("TableCell", parent=styles["BodyText"], fontName=pdf_font, fontSize=9, leading=11, wordWrap="CJK")
    table_header = ParagraphStyle("TableHeader", parent=styles["BodyText"], fontName=pdf_font, fontSize=9, leading=11, wordWrap="CJK")
    heading3 = ParagraphStyle("Heading3Pdf", parent=styles["Heading3"], fontName=pdf_font)
    heading4 = ParagraphStyle("Heading4Pdf", parent=styles["Heading4"], fontName=pdf_font)
    heading5 = ParagraphStyle("Heading5Pdf", parent=styles["Heading5"], fontName=pdf_font)
    heading6 = ParagraphStyle("Heading6Pdf", parent=styles["Heading6"], fontName=pdf_font)

    def pdf_cell(value: str, style=table_cell):
        return Paragraph(_escape_pdf_text(_normalize_text(value)), style)

    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story: List = []

    story.append(Paragraph(title, heading1))
    story.append(Paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}", small))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Attacker Profile", heading2))
    story.append(Paragraph(attacker_description, body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Questionnaire Results", heading2))
    story.append(Paragraph("The following section summarizes the completed questionnaires.", body))
    story.append(Spacer(1, 8))

    for label, groups in (("Technical Measures", technical_groups), ("Organizational Measures", organizational_groups)):
        if not groups:
            continue
        story.append(Paragraph(label, heading3))
        for summary in _group_summaries(groups):
            group_title = summary.get("group_name") or "Group"
            normalized = summary.get("normalized")
            normalized_text = f"{normalized:.2f}" if normalized is not None else "N/A"
            story.append(Paragraph(f"{group_title} (score: {normalized_text})", body))
            rows = [[
                pdf_cell("Question ID", table_header),
                pdf_cell("Question", table_header),
                pdf_cell("Score", table_header),
                pdf_cell("Max", table_header),
            ]]
            for q in summary.get("questions", []):
                rows.append(
                    [
                        pdf_cell(str(q.get("id", ""))),
                        pdf_cell(q.get("text")),
                        pdf_cell(_format_number(q.get("score"))),
                        pdf_cell(_format_number(q.get("max_score"))),
                    ]
                )
            table = Table(rows, colWidths=[60, 280, 60, 60])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 10))

    story.append(Paragraph("Baseline LOPA (Before Cyberattack)", heading2))
    story.append(
        Paragraph(
            "This diagram represents the protection layers and their base probabilities of failure on demand.",
            body,
        )
    )
    if lopa_before_path and os.path.exists(lopa_before_path):
        story.append(PdfImage(lopa_before_path, width=480, height=270))
    if layers_input:
        rows = [[
            pdf_cell("Layer", table_header),
            pdf_cell("Base PFD", table_header),
            pdf_cell("Cyber-prone", table_header),
        ]]
        for layer in layers_input:
            rows.append(
                [
                    pdf_cell(layer.get("name")),
                    pdf_cell(_format_number(layer.get("pfd"))),
                    pdf_cell("Yes" if layer.get("cyber", True) else "No"),
                ]
            )
        table = Table(rows, colWidths=[240, 100, 120])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ]
            )
        )
        story.append(table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Degraded LOPA (After Cyberattack)", heading2))
    story.append(
        Paragraph(
            "This diagram accounts for cyber degradation and organizational factors in the effective protection levels.",
            body,
        )
    )
    if lopa_after_path and os.path.exists(lopa_after_path):
        story.append(PdfImage(lopa_after_path, width=480, height=270))
    story.append(Spacer(1, 10))

    maturity_score = float(risk_result.get("maturity_score", 0.0))
    story.append(Paragraph("Maturity Summary", heading2))
    story.append(
        Paragraph(
            f"Maturity score: {_format_number(maturity_score)} ({_maturity_level(maturity_score)}).",
            body,
        )
    )
    story.append(Spacer(1, 10))

    story.append(Paragraph("Layer Degradation", heading2))
    layer_rows = [[
        pdf_cell("Layer", table_header),
        pdf_cell("Base PFD", table_header),
        pdf_cell("CDF", table_header),
        pdf_cell("ODF", table_header),
        pdf_cell("Effective PFD Y1", table_header),
        pdf_cell("Effective PFD Y2", table_header),
        pdf_cell("Effective PFD Y3", table_header),
    ]]
    for layer in risk_result.get("layers", []):
        layer_rows.append(
            [
                pdf_cell(layer.get("name")),
                pdf_cell(_format_number(layer.get("base_pfd"))),
                pdf_cell(_format_number(layer.get("degradation_factor"))),
                pdf_cell(_format_number(layer.get("org_multiplier"))),
                pdf_cell(_format_number(layer.get("effective_pfd"))),
                pdf_cell(_format_number(layer.get("effective_pfd_year2"))),
                pdf_cell(_format_number(layer.get("effective_pfd_year3"))),
            ]
        )
    layer_table = Table(layer_rows, colWidths=[120, 60, 60, 60, 70, 70, 70])
    layer_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(layer_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Event Losses", heading2))
    loss_rows = [[
        pdf_cell("Event", table_header),
        pdf_cell("Prob Y1", table_header),
        pdf_cell("Prob Y2", table_header),
        pdf_cell("Prob Y3", table_header),
        pdf_cell("SLE", table_header),
        pdf_cell("Currency", table_header),
        pdf_cell("Loss Y1", table_header),
        pdf_cell("Loss Y2", table_header),
        pdf_cell("Loss Y3", table_header),
    ]]
    for item in risk_result.get("event_losses", []):
        loss_rows.append(
            [
                pdf_cell(item.get("name")),
                pdf_cell(_format_number(item.get("probability_year1"))),
                pdf_cell(_format_number(item.get("probability_year2"))),
                pdf_cell(_format_number(item.get("probability_year3"))),
                pdf_cell(_format_number(item.get("sle"))),
                pdf_cell(item.get("currency")),
                pdf_cell(_format_number(item.get("loss_year1"))),
                pdf_cell(_format_number(item.get("loss_year2"))),
                pdf_cell(_format_number(item.get("loss_year3"))),
            ]
        )
    loss_table = Table(loss_rows, colWidths=[90, 50, 50, 50, 50, 50, 60, 60, 60])
    loss_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(loss_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Conclusions", heading2))
    rec_lines = _parse_markdown_lines(_normalize_text(recommendations_text))
    if not rec_lines:
        story.append(Paragraph("No recommendations provided.", body))
    else:
        heading_map = {
            1: heading1,
            2: heading2,
            3: heading3,
            4: heading4,
            5: heading5,
            6: heading6,
        }
        for level, content in rec_lines:
            if level is None:
                story.append(Spacer(1, 8))
                continue
            if level and level > 0:
                text = _escape_pdf_text(content)
                story.append(Paragraph(text, heading_map.get(level, heading2)))
                continue
            runs = _split_bold_runs(content)
            if not runs:
                story.append(Paragraph(_escape_pdf_text(content), body))
                continue
            rendered = "".join(
                f"<b>{_escape_pdf_text(text)}</b>" if is_bold else _escape_pdf_text(text)
                for text, is_bold in runs
            )
            story.append(Paragraph(rendered, body))

    doc.build(story)


def _build_docx(
    output_path: str,
    *,
    title: str,
    attacker_description: str,
    technical_groups: List[Dict],
    organizational_groups: List[Dict],
    layers_input: List[Dict],
    risk_result: Dict,
    lopa_before_path: Optional[str],
    lopa_after_path: Optional[str],
    recommendations_text: str,
) -> None:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_heading("Attacker Profile", level=1)
    doc.add_paragraph(attacker_description)

    doc.add_heading("Questionnaire Results", level=1)
    doc.add_paragraph("The following section summarizes the completed questionnaires.")

    for label, groups in (("Technical Measures", technical_groups), ("Organizational Measures", organizational_groups)):
        if not groups:
            continue
        doc.add_heading(label, level=2)
        for summary in _group_summaries(groups):
            group_title = summary.get("group_name") or "Group"
            normalized = summary.get("normalized")
            normalized_text = f"{normalized:.2f}" if normalized is not None else "N/A"
            doc.add_paragraph(f"{group_title} (score: {normalized_text})")
            table = doc.add_table(rows=1, cols=4)
            header_cells = table.rows[0].cells
            header_cells[0].text = "Question ID"
            header_cells[1].text = "Question"
            header_cells[2].text = "Score"
            header_cells[3].text = "Max"
            for q in summary.get("questions", []):
                row_cells = table.add_row().cells
                row_cells[0].text = _normalize_text(str(q.get("id", "")))
                row_cells[1].text = _normalize_text(q.get("text"))
                row_cells[2].text = _format_number(q.get("score"))
                row_cells[3].text = _format_number(q.get("max_score"))

    doc.add_heading("Baseline LOPA (Before Cyberattack)", level=1)
    doc.add_paragraph(
        "This diagram represents the protection layers and their base probabilities of failure on demand."
    )
    if lopa_before_path and os.path.exists(lopa_before_path):
        doc.add_picture(lopa_before_path, width=Inches(6.0))
    if layers_input:
        table = doc.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Layer"
        header_cells[1].text = "Base PFD"
        header_cells[2].text = "Cyber-prone"
        for layer in layers_input:
            row_cells = table.add_row().cells
            row_cells[0].text = _normalize_text(layer.get("name"))
            row_cells[1].text = _format_number(layer.get("pfd"))
            row_cells[2].text = "Yes" if layer.get("cyber", True) else "No"

    doc.add_heading("Degraded LOPA (After Cyberattack)", level=1)
    doc.add_paragraph(
        "This diagram accounts for cyber degradation and organizational factors in the effective protection levels."
    )
    if lopa_after_path and os.path.exists(lopa_after_path):
        doc.add_picture(lopa_after_path, width=Inches(6.0))

    maturity_score = float(risk_result.get("maturity_score", 0.0))
    doc.add_heading("Maturity Summary", level=1)
    doc.add_paragraph(f"Maturity score: {_format_number(maturity_score)} ({_maturity_level(maturity_score)}).")

    doc.add_heading("Layer Degradation", level=1)
    layer_table = doc.add_table(rows=1, cols=7)
    header_cells = layer_table.rows[0].cells
    header_cells[0].text = "Layer"
    header_cells[1].text = "Base PFD"
    header_cells[2].text = "CDF"
    header_cells[3].text = "ODF"
    header_cells[4].text = "Effective PFD Y1"
    header_cells[5].text = "Effective PFD Y2"
    header_cells[6].text = "Effective PFD Y3"
    for layer in risk_result.get("layers", []):
        row_cells = layer_table.add_row().cells
        row_cells[0].text = _normalize_text(layer.get("name"))
        row_cells[1].text = _format_number(layer.get("base_pfd"))
        row_cells[2].text = _format_number(layer.get("degradation_factor"))
        row_cells[3].text = _format_number(layer.get("org_multiplier"))
        row_cells[4].text = _format_number(layer.get("effective_pfd"))
        row_cells[5].text = _format_number(layer.get("effective_pfd_year2"))
        row_cells[6].text = _format_number(layer.get("effective_pfd_year3"))

    doc.add_heading("Event Losses", level=1)
    loss_table = doc.add_table(rows=1, cols=9)
    loss_header = loss_table.rows[0].cells
    loss_header[0].text = "Event"
    loss_header[1].text = "Prob Y1"
    loss_header[2].text = "Prob Y2"
    loss_header[3].text = "Prob Y3"
    loss_header[4].text = "SLE"
    loss_header[5].text = "Currency"
    loss_header[6].text = "Loss Y1"
    loss_header[7].text = "Loss Y2"
    loss_header[8].text = "Loss Y3"
    for item in risk_result.get("event_losses", []):
        row_cells = loss_table.add_row().cells
        row_cells[0].text = _normalize_text(item.get("name"))
        row_cells[1].text = _format_number(item.get("probability_year1"))
        row_cells[2].text = _format_number(item.get("probability_year2"))
        row_cells[3].text = _format_number(item.get("probability_year3"))
        row_cells[4].text = _format_number(item.get("sle"))
        row_cells[5].text = _normalize_text(item.get("currency"))
        row_cells[6].text = _format_number(item.get("loss_year1"))
        row_cells[7].text = _format_number(item.get("loss_year2"))
        row_cells[8].text = _format_number(item.get("loss_year3"))

    doc.add_heading("Conclusions", level=1)
    rec_lines = _parse_markdown_lines(_normalize_text(recommendations_text))
    if not rec_lines:
        doc.add_paragraph("No recommendations provided.")
    else:
        for level, content in rec_lines:
            if level is None:
                doc.add_paragraph("")
                continue
            if level and level > 0:
                doc.add_heading(content, level=min(level, 6))
                continue
            paragraph = doc.add_paragraph("")
            runs = _split_bold_runs(content)
            if not runs:
                paragraph.add_run(content)
                continue
            for text, is_bold in runs:
                run = paragraph.add_run(text)
                run.bold = is_bold

    doc.save(output_path)


def generate_report(
    *,
    report_format: str,
    technical_groups: List[Dict],
    organizational_groups: List[Dict],
    layers_input: List[Dict],
    risk_result: Dict,
    attacker_type: str,
    attacker_potential: str,
    recommendations_text: str,
) -> str:
    report_format = (report_format or "pdf").lower()
    output_dir = tempfile.mkdtemp(prefix="ics_report_")
    layers_after = [
        {"name": layer.get("name"), "pfd": layer.get("effective_pfd", 1.0)} for layer in risk_result.get("layers", [])
    ]
    lopa_before_path, lopa_after_path = _build_lopa_images(layers_input, layers_after, output_dir)

    title = "ICS Cybersecurity Risk Assessment Report"
    attacker_description = _format_attacker(attacker_type, attacker_potential)

    if report_format == "docx":
        output_path = os.path.join(output_dir, "risk_report.docx")
        _build_docx(
            output_path,
            title=title,
            attacker_description=attacker_description,
            technical_groups=technical_groups,
            organizational_groups=organizational_groups,
            layers_input=layers_input,
            risk_result=risk_result,
            lopa_before_path=lopa_before_path,
            lopa_after_path=lopa_after_path,
            recommendations_text=recommendations_text,
        )
        return output_path

    output_path = os.path.join(output_dir, "risk_report.pdf")
    _build_pdf(
        output_path,
        title=title,
        attacker_description=attacker_description,
        technical_groups=technical_groups,
        organizational_groups=organizational_groups,
        layers_input=layers_input,
        risk_result=risk_result,
        lopa_before_path=lopa_before_path,
        lopa_after_path=lopa_after_path,
        recommendations_text=recommendations_text,
    )
    return output_path
